from sqlalchemy import create_engine, inspect
import pandas as pd
import re
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
import warnings

PATTERNS = {
    'email': re.compile(r'^[^@]+@[^@]+\.[^@]+$'),
    'phone': re.compile(r'^1[3-9]\d{9}$'),
    'date': re.compile(r'^\d{4}-\d{2}-\d{2}$'),
    # 可扩展
}

def guess_best_pattern(values):
    samples = values.dropna().astype(str)
    if len(samples) == 0:
        return None, None
    samples = samples.sample(min(20, len(samples)), random_state=42)
    best_pat, best_rate = None, 0
    for name, pat in PATTERNS.items():
        match_count = sum(bool(pat.match(v)) for v in samples)
        rate = match_count / len(samples)
        if rate > best_rate:
            best_pat, best_rate = pat, rate
    if best_rate > 0.8:
        return best_pat, best_rate
    return None, None

def guess_field_type(values):
    import re
    # 只取非空样本
    samples = values.dropna().astype(str)
    if len(samples) == 0:
        return 'unknown'
    samples = samples.sample(min(20, len(samples)), random_state=42)
    patterns = {
        'email': re.compile(r'^[^@]+@[^@]+\.[^@]+$'),
        'phone': re.compile(r'^1[3-9]\d{9}$'),
        'date': re.compile(r'^\d{4}-\d{2}-\d{2}'),
        # 可扩展更多
    }
    match_counts = {k: 0 for k in patterns}
    for v in samples:
        for k, pat in patterns.items():
            if pat.match(v):
                match_counts[k] += 1
    for k, cnt in match_counts.items():
        if cnt / len(samples) > 0.8:
            return k
    # 尝试是否为数值
    try:
        pd.to_numeric(samples)
        return 'numeric'
    except Exception:
        pass
    # 优化日期推断，先尝试常见格式，最后才fallback
    date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%Y-%m-%d %H:%M:%S', '%Y/%m/%d %H:%M:%S']
    for fmt in date_formats:
        try:
            pd.to_datetime(samples, format=fmt)
            return 'date'
        except Exception:
            continue
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            pd.to_datetime(samples)  # fallback, 彻底静默警告
        return 'date'
    except Exception:
        pass
    return 'string'


def profiling_and_report(db_path, report_name, ignore_table_list=None, ignore_column_dict=None):
    if ignore_table_list is None:
        ignore_table_list = ['transactions']
    if ignore_column_dict is None:
        ignore_column_dict = {'user_info': ['region']}
    engine = create_engine(db_path)
    inspector = inspect(engine)
    all_tables = inspector.get_table_names()
    report = {}
    report['tables'] = {}
    report['missing'] = {}
    report['abnormal'] = {}
    report['duplicate'] = {}
    report['distribution'] = {}
    report['highlight'] = {}  # 新增高亮map

    for table in all_tables:
        df = pd.read_sql(table, engine)
        report['tables'][table] = list(df.columns)
        # 缺失值检测
        for col in df.columns:
            missing = df[df[col].isnull()]
            if not missing.empty:
                # 保存整行和高亮map
                key = f'{table}.{col}'
                report['missing'][key] = df.loc[missing.index]
                report['highlight'][key] = {(i, col): True for i in missing.index}
        # 字段类型推断
        for col in df.columns:
            values = df[col]
            best_pat, best_rate = guess_best_pattern(values)
            if best_pat is not None:
                abnormal = df[~values.astype(str).str.match(best_pat, na=False)]
                if not abnormal.empty:
                    key = f'{table}.{col}_pattern'
                    report['abnormal'][key] = df.loc[abnormal.index]
                    report['highlight'][key] = {(i, col): True for i in abnormal.index}
            # 数值型异常
            elif pd.api.types.is_numeric_dtype(values):
                mean = values.mean()
                std = values.std()
                if std > 0:
                    abnormal = df[(values < mean - 5*std) | (values > mean + 5*std)]
                    if not abnormal.empty:
                        key = f'{table}.{col}'
                        report['abnormal'][key] = df.loc[abnormal.index]
                        report['highlight'][key] = {(i, col): True for i in abnormal.index}
            # 字符串型异常
            elif pd.api.types.is_string_dtype(values):
                too_short = df[values.astype(str).str.len() < 1]
                too_long = df[values.astype(str).str.len() > 50]
                if not too_short.empty:
                    key = f'{table}.{col}_too_short'
                    report['abnormal'][key] = df.loc[too_short.index]
                    report['highlight'][key] = {(i, col): True for i in too_short.index}
                if not too_long.empty:
                    key = f'{table}.{col}_too_long'
                    report['abnormal'][key] = df.loc[too_long.index]
                    report['highlight'][key] = {(i, col): True for i in too_long.index}
        # 重复值检测
        if table not in ignore_table_list:
            for col in df.columns:
                if table in ignore_column_dict and col in ignore_column_dict[table]:
                    continue
                duplicate = df[df.duplicated([col], keep=False)]
                if not duplicate.empty:
                    key = f'{table}.{col}'
                    report['duplicate'][key] = df.loc[duplicate.index]
                    report['highlight'][key] = {(i, col): True for i in duplicate.index}

    # 生成docx报告
    doc = Document()
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    doc.add_heading('数据源Profiling报告', 0)
    doc.add_heading('1. 表结构', level=1)
    for table, cols in report['tables'].items():
        p = doc.add_paragraph(f"{table}: {', '.join(cols)}")
        for run in p.runs:
            set_font(run)

    doc.add_heading('2. 缺失值检测', level=1)
    for key, df in report['missing'].items():
        doc.add_paragraph(f"{key} 缺失值记录数: {len(df)}")
        add_df_table(doc, df, highlight_map=report['highlight'].get(key))

    doc.add_heading('3. 异常值检测', level=1)
    for key, df in report['abnormal'].items():
        doc.add_paragraph(f"{key} 异常记录数: {len(df)}")
        add_df_table(doc, df, highlight_map=report['highlight'].get(key))

    doc.add_heading('4. 重复值检测', level=1)
    for key, df in report['duplicate'].items():
        doc.add_paragraph(f"{key} 重复记录数: {len(df)}")
        add_df_table(doc, df, highlight_map=report['highlight'].get(key))

    doc.save(report_name)
    print("报告已生成："+report_name)
    return report

def set_font(run, font_name='微软雅黑', font_size=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size

def add_df_table(doc, df, max_rows=100, highlight_map=None):
    if df.empty:
        doc.add_paragraph("无记录")
        return
    df = df.head(max_rows)
    table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        run = hdr_cells[i].paragraphs[0].add_run(str(col))
        set_font(run)
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            run = row_cells[i].paragraphs[0].add_run(str(row[col]))
            set_font(run)
            if highlight_map and (idx, col) in highlight_map:
                # 标红底色
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                row_cells[i]._tc.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="FF6666"/>'.format(nsdecls('w')))
                )