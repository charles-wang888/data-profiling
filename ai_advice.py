import requests
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn

from profiling_report import set_font

OLLAMA_API_URL = 'http://localhost:11434/api/generate'
OLLAMA_MODEL = 'qwen3:4b'

def ollama_generate(prompt, model=OLLAMA_MODEL):
    headers = {"Content-Type": "application/json"}
    data = {
        "model": model,
        "prompt": prompt,
        "stream": False
    }
    try:
        response = requests.post(OLLAMA_API_URL, json=data, headers=headers, timeout=60)
        response.raise_for_status()
        return response.json().get('response', '')
    except Exception as e:
        print("Ollama API调用失败：", e)
        return "Ollama API调用失败"



def ai_advice_and_fix(report, report_name):
    # 动态生成字段信息
    field_info = []
    for table, cols in report['tables'].items():
        field_info.append(f"{table}: {', '.join(cols)}")
    field_info_str = '\n'.join(field_info)

    # 校验建议Prompt
    check_prompt = (
        f"以下是数据表的字段信息：\n{field_info_str}\n"
        "请你根据字段名和常见数据质量问题，自动给出每个字段的校验建议。"
        "请先用<think><think>写一段AI的思考过程，再用表格形式给出详细建议。"
    )
    check_advice = ollama_generate(check_prompt)
    check_think, check_table = check_advice.split('</think>') if '</think>' in check_advice else (check_advice, '')

    # 动态生成异常/缺失/重复数据字符串
    def get_problem_str(report, key):
        s = []
        for field, df in report[key].items():
            if hasattr(df, 'to_string'):
                s.append(f"{field}:\n{df.to_string(index=False)[:1000]}")
        return '\n'.join(s)

    # 修复建议Prompt
    fix_prompt = (
        f"请根据以下异常数据，给出修复建议：\n"
        f"1. 格式错误：建议用正确正则修正，列出'值-建议值'\n"
        f"2. 缺失值：是否有默认值补全，列出修复建议\n"
        f"3. 重复值：建议保留最新一条，列出重复记录\n"
        f"请先用<think><think>写一段AI的思考过程，再用表格形式给出详细建议。\n"
        f"异常值：\n{get_problem_str(report, 'abnormal')}\n"
        f"缺失值：\n{get_problem_str(report, 'missing')}\n"
        f"重复值：\n{get_problem_str(report, 'duplicate')}\n"
    )
    fix_advice = ollama_generate(fix_prompt)
    fix_think, fix_table = fix_advice.split('</think>') if '</think>' in fix_advice else (fix_advice, '')

    # 写入docx报告
    doc = Document(report_name)
    doc.styles['Normal'].font.name = u'微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    # 校验建议
    add_think_box(doc, 'AI校验建议思考过程', check_think.replace('<think>', '').strip())
    add_advice_table(doc, check_table.strip())

    # 修复建议
    add_think_box(doc, 'AI修复建议思考过程', fix_think.replace('<think>', '').strip())
    add_advice_table(doc, fix_table.strip())

    doc.save(report_name)
    print("AI建议已写入报告。")

def add_think_box(doc, title, think_text):
    # 添加标题
    doc.add_paragraph(title, style='Heading 2')
    # 创建1行1列的表格
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = think_text
    # 设置底色为浅灰色
    cell._tc.get_or_add_tcPr().append(
        parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
    )
    # 设置字体
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_font(run)

def add_advice_table(doc, advice_text):
    # 按行分割，填入表格
    lines = [line for line in advice_text.split('\n') if line.strip()]
    table = doc.add_table(rows=len(lines), cols=1, style='Table Grid')
    for i, line in enumerate(lines):
        run = table.cell(i, 0).paragraphs[0].add_run(line)
        set_font(run)